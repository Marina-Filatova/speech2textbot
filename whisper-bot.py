from docx import Document
from recognizer import Recognizer
import subprocess
import telebot
import os 
import speech_recognition as sr
import datetime
import requests
import soundfile as sf 
from aiogram import Bot, Dispatcher , types

logfile = str(datetime.date.today()) + '.log' # формируем имя лог-файла
token = 'your_bot_token'
bot = telebot.TeleBot(token)

# Define the keyboard layout
keyboard = telebot.types.ReplyKeyboardMarkup(row_width=2, resize_keyboard=True)
start_button = telebot.types.KeyboardButton('Привет!')
select_output_button = telebot.types.KeyboardButton('Выбрать формат вывода')
keyboard.add(start_button, select_output_button)

# Define the output format keyboard layout
output_format_keyboard = telebot.types.ReplyKeyboardMarkup(row_width=3, resize_keyboard=True)
txt_button = telebot.types.KeyboardButton('.txt')
doc_button = telebot.types.KeyboardButton('.doc')
message_button = telebot.types.KeyboardButton('Просто отправить сообщением')
output_format_keyboard.add(txt_button, doc_button, message_button)

selected_format = None

@bot.message_handler(commands=['start'])
def start_command(message):
    bot.send_message(message.chat.id, f"Приятно познакомиться, {message.from_user.first_name}. Я бот, преобразующий аудио в текст. При работе со мной тебе стоит помнить, что я не очень быстр.", reply_markup=keyboard)

@bot.message_handler(content_types=['text'])
def get_messages(message):
    global selected_format
    if message.text == "Привет!":
        bot.send_message(message.from_user.id, "Привет! Я - бот, преобразующий аудио в текст. Чем я могу помочь?", reply_markup=keyboard)
    # elif message.text == "Перезапуск":
    elif message.text == "/help":
        bot.send_message(message.from_user.id, "1. Чтобы начать работу с ботом напиши в чат команду /start \n2. Чтобы получить результат в формате .txt - нажми в меню кнопку «выбрать формат вывода» после чего выбери кнопку .txt\n3. Чтобы получить результат в формате .docx - нажми в меню кнопку «выбрать формат вывода» после чего выбери кнопку .doc \n4. Если хочешь получить обычное сообщение в чат - нажми в меню кнопку «выбрать формат вывода» после чего выбери кнопку «отправить сообщение в чат».\n5. Для того, чтобы перезапустить бота и очистить историю, нажми кнопку «Перезапуск» в меню \n!!Бот работает как с голосовыми, так и с файлами в формате .mp3, просто отправь боту файл со своего устройства или запиши голосовое в чат!!", reply_markup=keyboard)
    # elif message.text == "Start":
    #     bot.send_message(message.from_user.id, "Let's get started!", reply_markup=keyboard)
    elif message.text == "Выбрать формат вывода":
        bot.send_message(message.from_user.id, "Пожалуйста выберите формат:", reply_markup=output_format_keyboard)
    elif message.text in ['.txt', '.doc', 'Просто отправить сообщением']:
        selected_format = message.text
        bot.send_message(message.from_user.id, f"Вы выбрали формат: {selected_format}. Пожалуйста пришлите Ваше сообщение. Я могу работать с голосовыми сообщениями и аудиофайлами.", reply_markup=keyboard)
    else:
        bot.send_message(message.from_user.id, "Извините, я не понимаю. Для получения инструкций отправь мне /help.", reply_markup=keyboard)

# Handle the output format selection
@bot.message_handler(content_types=['voice'])
def get_audio_messages(message):
    global selected_format
    try:
        bot.send_message(message.from_user.id, "Подождите, я в ускоренном темпе слушаю Ваше голосовое. Скоро я отправлю результат")
        print("Started recognition...")
        # Extract the file name from the message
        file_info = bot.get_file(message.voice.file_id)
        fname = os.path.basename(file_info.file_path)
        # Download and save the audio file
        doc = requests.get('https://api.telegram.org/file/bot{0}/{1}'.format(token, file_info.file_path))
        with open(fname, 'wb') as f:
            f.write(doc.content)

        # Convert the audio file to .mp3 if it's not already in that format
        if not fname.endswith('.mp3'):
            data, samplerate = sf.read(fname)
            sf.write(fname + '.mp3', data, samplerate)
            os.remove(fname)
            fname += '.mp3'

        with open(fname, "rb") as f:
            audio_bin = f.read()

        with Recognizer() as recognizer:
            result = recognizer.recognize(audio_bin)

        print("Recognition completed!")

        # Send the result based on the selected format
        if selected_format == '.txt':
            with open(f"{message.from_user.id}.txt", "w", encoding='utf-8') as f:
                f.write(result)
            with open(f"{message.from_user.id}.txt", "rb") as f:
                bot.send_document(message.from_user.id, f)
        elif selected_format == '.doc':
            # Implement .doc conversion here
            doc = Document()
            doc.add_paragraph(result)
            doc_filename = f"{message.from_user.id}.docx"
            doc.save(doc_filename)
            os.remove(doc_filename)
            # Send the .docx file to the user
            with open(doc_filename, "rb") as doc_file:
                bot.send_document(message.from_user.id, doc_file)
            # Remove the temporary .docx file
            os.remove(doc_filename)
        else:
            bot.send_message(message.from_user.id, result)

        # Reset the selected format
        selected_format = None

    except sr.UnknownValueError as e:
        bot.send_message(message.from_user.id, "Прошу прощения, но я не разобрал сообщение, или оно не содержит важной информации...")
        with open(logfile, 'a', encoding='utf-8') as f:
            f.write(str(datetime.datetime.today().strftime("%H:%M:%S")) + ':' + str(message.from_user.id) + ': Сообщение пустое.\n')
    except Exception as e:
        bot.send_message(message.from_user.id, "Что-то пошло не так, мы уже трудимся над решением этой проблемы")
        with open(logfile, 'a', encoding='utf-8') as f:
            f.write(str(datetime.datetime.today().strftime("%H:%M:%S")) + ':' + str(message.from_user.id) + ':' + str(e) + '\n')
    finally:
        # Remove the downloaded files
        if os.path.exists(fname):
            os.remove(fname)

# функция для обработки аудиофайлов
@bot.message_handler(content_types=['audio'])
def get_audio_file(message):
    global selected_format
    try:
        bot.send_message(message.from_user.id, "Подождите, я в ускоренном темпе слушаю Ваш файл. Скоро я отправлю результат")
        print("Started recognition...")
        # Ниже пытаемся вычленить имя файла из сообщения
        file_info = bot.get_file(message.audio.file_id)
        # полный путь до файла (например: music/file_2.m4a)
        path = os.path.splitext(file_info.file_path)[0]
        only_audio_name = os.path.basename(path)
        # Преобразуем путь в имя файла (например: file_2.m4a)
        fname = os.path.basename(file_info.file_path)
        print(fname)

        # Получаем и сохраняем файл
        doc = requests.get('https://api.telegram.org/file/bot{0}/{1}'.format(token, file_info.file_path)) 
        with open(fname, 'wb') as f:
            # сохранение аудио-сообщения
            f.write(doc.content) 
        
        # преобразование в нужный нам формат
        process = subprocess.run(['ffmpeg', '-i', fname, only_audio_name+'_output.mp3'])# здесь используется ffmpeg, для конвертации в .mp3

        with Recognizer() as recognizer:
            result = recognizer.recognize(only_audio_name+'_output.mp3')

        print("Recognition completed!")
       
        # Send the result based on the selected format
        if selected_format == '.txt':
            with open(f"{message.from_user.id}.txt", "w", encoding='utf-8') as f:
                f.write(result)
            with open(f"{message.from_user.id}.txt", "rb") as f:
                bot.send_document(message.from_user.id, f)
        elif selected_format == '.doc':
            # Implement .doc conversion here
            doc = Document()
            doc.add_paragraph(result)
            doc_filename = f"{message.from_user.id}.docx"
            doc.save(doc_filename)
            os.remove(doc_filename)
            # Send the .docx file to the user
            with open(doc_filename, "rb") as doc_file:
                bot.send_document(message.from_user.id, doc_file)
            # Remove the temporary .docx file
            os.remove(doc_filename)
        else:
            bot.send_message(message.from_user.id, result)

        # Reset the selected format
        selected_format = None
    
    except sr.UnknownValueError as e:
        bot.send_message(message.from_user.id,  "Прошу прощения, но я не разобрал сообщение, или оно поустое...")
        with open(logfile, 'a', encoding='utf-8') as f:
            f.write(str(datetime.datetime.today().strftime("%H:%M:%S")) + ':' + str(message.from_user.id) + ':' + str(message.from_user.first_name) + '_' + str(message.from_user.last_name) + ':' + str(message.from_user.username) +':'+ str(message.from_user.language_code) + ': Сообщение пустое.\n')
    except Exception as e:
        bot.send_message(message.from_user.id,  "Что-то пошло пошло не так, мы уже трудимся над решением этой проблемы")
        with open(logfile, 'a', encoding='utf-8') as f:
            f.write(str(datetime.datetime.today().strftime("%H:%M:%S")) + ':' + str(message.from_user.id) + ':' + str(message.from_user.first_name) + '_' + str(message.from_user.last_name) + ':' + str(message.from_user.username) +':'+ str(message.from_user.language_code) +':' + str(e) + '\n')
    finally:
        # Remove the downloaded files
        if os.path.exists(fname):
            os.remove(fname)  
        if os.path.exists(only_audio_name+'_output.mp3'):
            os.remove(only_audio_name+'_output.mp3')
        if os.path.exists(only_audio_name+'.mp3'):
            os.remove(only_audio_name+'.mp3')
        

# Run the bot
bot.polling(none_stop=True)